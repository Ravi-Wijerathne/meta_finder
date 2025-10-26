"""Document metadata extraction."""
import os

try:
    from PyPDF2 import PdfReader
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


def extract(file_path):
    """
    Extract metadata from document files.
    
    Args:
        file_path: Path to document file
        
    Returns:
        dict: Metadata dictionary
    """
    metadata = {}
    
    # Add basic file info
    metadata.update(get_basic_info(file_path))
    
    # Determine document type and extract
    ext = os.path.splitext(file_path)[1].lower()
    
    if ext == '.pdf' and PYPDF2_AVAILABLE:
        metadata.update(extract_pdf(file_path))
    elif ext in ['.docx', '.doc'] and DOCX_AVAILABLE:
        metadata.update(extract_docx(file_path))
    elif ext == '.txt':
        metadata.update(extract_txt(file_path))
    else:
        metadata['note'] = f"No specific extractor for {ext} files"
    
    return metadata


def get_basic_info(file_path):
    """Get basic file information."""
    stat = os.stat(file_path)
    return {
        'file_name': os.path.basename(file_path),
        'file_size_bytes': stat.st_size,
        'file_size_mb': round(stat.st_size / (1024 * 1024), 2),
    }


def extract_pdf(file_path):
    """Extract PDF metadata."""
    metadata = {}
    
    try:
        reader = PdfReader(file_path)
        
        # Document info
        metadata['num_pages'] = len(reader.pages)
        
        # Metadata
        if reader.metadata:
            meta = reader.metadata
            metadata['pdf_title'] = meta.get('/Title', '')
            metadata['pdf_author'] = meta.get('/Author', '')
            metadata['pdf_subject'] = meta.get('/Subject', '')
            metadata['pdf_creator'] = meta.get('/Creator', '')
            metadata['pdf_producer'] = meta.get('/Producer', '')
            metadata['pdf_creation_date'] = meta.get('/CreationDate', '')
            metadata['pdf_modification_date'] = meta.get('/ModDate', '')
            
            # Add any other metadata fields
            for key, value in meta.items():
                if key not in ['/Title', '/Author', '/Subject', '/Creator', 
                              '/Producer', '/CreationDate', '/ModDate']:
                    clean_key = key.strip('/').replace('/', '_')
                    metadata[f'pdf_{clean_key}'] = str(value)
        
        # Check if encrypted
        metadata['is_encrypted'] = reader.is_encrypted
        
        # Extract text from first page (preview)
        if len(reader.pages) > 0:
            first_page_text = reader.pages[0].extract_text()
            preview = first_page_text[:200] if first_page_text else ""
            metadata['first_page_preview'] = preview.replace('\n', ' ')
    
    except Exception as e:
        metadata['pdf_error'] = str(e)
    
    return metadata


def extract_docx(file_path):
    """Extract DOCX metadata."""
    metadata = {}
    
    try:
        doc = Document(file_path)
        
        # Core properties
        props = doc.core_properties
        metadata['docx_title'] = props.title or ''
        metadata['docx_author'] = props.author or ''
        metadata['docx_subject'] = props.subject or ''
        metadata['docx_keywords'] = props.keywords or ''
        metadata['docx_comments'] = props.comments or ''
        metadata['docx_category'] = props.category or ''
        metadata['docx_created'] = str(props.created) if props.created else ''
        metadata['docx_modified'] = str(props.modified) if props.modified else ''
        metadata['docx_last_modified_by'] = props.last_modified_by or ''
        metadata['docx_revision'] = props.revision
        
        # Document statistics
        metadata['num_paragraphs'] = len(doc.paragraphs)
        metadata['num_tables'] = len(doc.tables)
        metadata['num_sections'] = len(doc.sections)
        
        # Extract first paragraph as preview
        if doc.paragraphs:
            first_para = doc.paragraphs[0].text
            metadata['first_paragraph_preview'] = first_para[:200]
    
    except Exception as e:
        metadata['docx_error'] = str(e)
    
    return metadata


def extract_txt(file_path):
    """Extract basic text file info."""
    metadata = {}
    
    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()
            
        metadata['num_characters'] = len(content)
        metadata['num_lines'] = content.count('\n') + 1
        metadata['num_words'] = len(content.split())
        
        # Preview
        preview = content[:200] if content else ""
        metadata['preview'] = preview.replace('\n', ' ')
    
    except Exception as e:
        metadata['txt_error'] = str(e)
    
    return metadata
